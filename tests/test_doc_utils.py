# -*- coding: utf-8 -*-
"""
针对 utils.doc_utils 的通用函数编写的 pytest 单元测试。
覆盖点：
- 段落格式应用（对齐、首行缩进、行距、段后）
- 标题级别计算（默认与自定义映射）
- 标题文本格式化（title 类型添加书名号样式）
- 返回目录：占位符创建与后续转换为指向书签的内部超链接
- 书签：在段落 XML 中插入 bookmarkStart/bookmarkEnd 而不改变文本
"""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from csv_word_converter.utils.doc_utils import (
    apply_paragraph_format,
    compute_heading_level,
    format_title_text,
    add_return_directory_placeholder,
    add_bookmark_to_paragraph_xml,
    create_target_bookmark_by_keyword,
    convert_return_placeholders_to_hyperlinks,
)


def test_apply_paragraph_format_alignment_and_spacing():
    """验证 apply_paragraph_format 能正确设置对齐、首行缩进、行距与段后。"""
    doc = Document()
    p = doc.add_paragraph("hello")
    cfg = {
        "alignment": "center",
        "first_line_indent": 2,  # 2 个字符，对应 Pt(24)
        "line_spacing": 1.5,
        "space_after": 12,
    }
    apply_paragraph_format(p, cfg)

    assert p.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert float(p.paragraph_format.first_line_indent.pt) == 24.0
    assert float(p.paragraph_format.line_spacing) == 1.5
    # space_after 以磅为单位
    assert float(p.paragraph_format.space_after.pt) == 12.0


def test_compute_heading_level_defaults_and_mapping():
    """验证 compute_heading_level 的默认映射与自定义映射行为。"""
    assert compute_heading_level("heading_1") == 1
    assert compute_heading_level("heading_2") == 2
    assert compute_heading_level("heading_3") == 3
    assert compute_heading_level("title") == 3
    assert compute_heading_level("unknown") is None

    custom = {"heading_1": 2, "heading_2": 3, "title": 1}
    assert compute_heading_level("heading_1", custom) == 2
    assert compute_heading_level("heading_2", custom) == 3
    assert compute_heading_level("title", custom) == 1


def test_format_title_text():
    """验证 title 类型文本会被【】包裹，其他类型保持原样。"""
    assert format_title_text("目录", "title") == "【目录】"
    assert format_title_text("第一章", "heading_1") == "第一章"


def test_add_bookmark_to_paragraph_xml():
    """为包含“目录”的段落插入书签，检查 bookmarkStart/bookmarkEnd 是否存在。"""
    doc = Document()
    p = doc.add_paragraph("目录")
    add_bookmark_to_paragraph_xml(p, "目录")

    # 使用 local-name()，避免 BaseOxmlElement.xpath 不支持 namespaces 参数的问题
    starts = p._element.xpath('.//*[local-name()="bookmarkStart"]')
    ends = p._element.xpath('.//*[local-name()="bookmarkEnd"]')

    assert len(starts) == 1
    assert len(ends) == 1
    # 名称写在 bookmarkStart 的 name 属性上
    assert starts[0].get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}name") == "目录"


def test_add_and_convert_return_placeholder_to_hyperlink():
    """
    1) 先在文档中定位“目录”并创建目标书签；
    2) 添加“返回目录”占位符段落；
    3) 将占位符转换为指向“目录”书签的内部超链接；
    4) 验证段落 XML 中产生 w:hyperlink 且锚点为“目录”。
    """
    doc = Document()
    # 目标段落与书签
    doc.add_paragraph("目录")
    assert create_target_bookmark_by_keyword(doc, "目录") is True

    # 添加占位符
    add_return_directory_placeholder(
        doc,
        {
            "text": "返回目录",
            "alignment": "right",
            "underline": True,
            "font_name": "宋体",
            "font_size": 12,
        },
    )

    # 转换为内部超链接
    convert_return_placeholders_to_hyperlinks(
        doc,
        placeholder_text="返回目录",
        bookmark_name="目录",
    )
    # 重构后函数不返回计数，断言转换成功与否依赖后续的XML结构检查

    # 验证超链接是否正确生成
    # 由于 python-docx 的 p.text 不会提取超链接内的文本，我们直接检查底层 XML
    target_paragraph = doc.paragraphs[-1]  # "返回目录" 是最后添加的段落

    # 1. 检查段落中是否包含 w:hyperlink 元素
    hyperlinks = target_paragraph._p.xpath('.//w:hyperlink')
    assert len(hyperlinks) >= 1, "未在段落中找到超链接元素"

    # 2. 检查超链接的锚点（书签）是否正确
    hyperlink_element = hyperlinks[0]
    anchor_attr = hyperlink_element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}anchor')
    assert anchor_attr == "目录", f"超链接锚点不正确，应为'目录'，实际为'{anchor_attr}'"

    # 3. 检查超链接显示的文本是否正确
    text_nodes = hyperlink_element.xpath('.//w:t/text()')
    hyperlink_text = "".join(text_nodes)
    assert hyperlink_text == "返回目录", f"超链接文本不正确，应为'返回目录'，实际为'{hyperlink_text}'"